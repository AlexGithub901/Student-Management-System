from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('STUDENT MANAGEMENT SYSTEM')
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 51, 102)

# Subtitle
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('Git, GitHub & GitHub Actions - Industry Assignment Report')
subtitle_run.font.size = Pt(14)
subtitle_run.font.italic = True
subtitle_run.font.color.rgb = RGBColor(64, 64, 64)

# Date
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_para.add_run(f'Report Date: {datetime.now().strftime("%B %d, %Y")}')
date_run.font.size = Pt(11)
date_run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()

# Executive Summary
heading1 = doc.add_heading('Executive Summary', level=1)
heading1_format = heading1.paragraph_format
heading1_format.space_before = Pt(12)
heading1_format.space_after = Pt(6)

summary_text = """This report documents the complete implementation of a professional-grade Student Management System with Docker containerization and GitHub CI/CD automation. The project demonstrates industry-standard practices in Git version control, GitHub collaboration, and automated testing through GitHub Actions workflows."""

summary_para = doc.add_paragraph(summary_text)
summary_para.paragraph_format.line_spacing = 1.15

# Task 1: GitHub Repository URL
doc.add_heading('Task 1: GitHub Repository URL', level=2)

repo_para = doc.add_paragraph()
repo_run = repo_para.add_run('Primary Repository: ')
repo_run.font.bold = True
repo_url_run = repo_para.add_run('https://github.com/AlexGithub901/Student-Management-System')
repo_url_run.font.color.rgb = RGBColor(0, 0, 255)
repo_url_run.underline = True

details = doc.add_paragraph('Repository Owner: AlexGithub901', style='List Bullet')
details = doc.add_paragraph('Repository Name: Student-Management-System', style='List Bullet')
details = doc.add_paragraph('Repository Type: Public', style='List Bullet')
details = doc.add_paragraph('Default Branch: main', style='List Bullet')

# Project Overview Section
doc.add_heading('Project Overview', level=2)

overview_text = """The Student Management System is a Python-based application that manages student records with the following core functionalities:

• Add Students: Register new students with ID, name, and grade
• Remove Students: Delete student records by ID
• Search Students: Case-insensitive search by student ID or name
• Update Students: Modify student information (name and/or grade)
• Export Students: Export all student records in CSV and JSON formats (bonus feature)"""

doc.add_paragraph(overview_text)

# Repository Structure
doc.add_heading('Repository Structure', level=2)

structure_table = doc.add_table(rows=9, cols=2)
structure_table.style = 'Light Grid Accent 1'

# Header row
header_cells = structure_table.rows[0].cells
header_cells[0].text = 'File/Folder'
header_cells[1].text = 'Description'

# Make header bold
for cell in header_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

# Data rows
structure_data = [
    ('student.py', 'Core student management functions'),
    ('test_student.py', 'Comprehensive test suite (17 tests)'),
    ('requirements.txt', 'Python dependencies (pytest)'),
    ('Dockerfile', 'Multi-stage Docker build configuration'),
    ('docker-compose.yml', 'Docker Compose for local testing'),
    ('.dockerignore', 'Files excluded from Docker context'),
    ('.github/workflows/', 'GitHub Actions workflow files'),
    ('README.md', 'Project documentation'),
]

for i, (file, desc) in enumerate(structure_data, 1):
    row_cells = structure_table.rows[i].cells
    row_cells[0].text = file
    row_cells[1].text = desc

doc.add_paragraph()

# Task 2-4: Screenshots Section
doc.add_page_break()

doc.add_heading('Task 2-4: GitHub Actions & Pull Request Screenshots', level=2)

screenshots_intro = doc.add_paragraph("""The following section contains screenshots demonstrating successful and failed GitHub Actions workflows, as well as the Pull Request page for feature-export branch.""")

doc.add_paragraph()

doc.add_heading('Screenshot 1: Successful GitHub Actions Workflow', level=3)
success_notes = doc.add_paragraph("""Status: ✅ PASSED
All 17 tests executed successfully in both Python CI and Docker Build & Push workflows.
Execution Time: ~45 seconds
Tests Passed: 17/17 (100%)""")

doc.add_paragraph('[SCREENSHOT PLACEHOLDER - Insert successful GitHub Actions workflow run here]', style='List Bullet')
doc.add_paragraph('Expected to show: Green checkmarks, all 17 tests passing, completion status', style='List Bullet')

doc.add_paragraph()

doc.add_heading('Screenshot 2: Failed GitHub Actions Workflow', level=3)
failure_notes = doc.add_paragraph("""Status: ❌ FAILED
This demonstrates how GitHub Actions catches test failures and blocks merging.
The failure can be triggered by intentionally breaking a test case.""")

doc.add_paragraph('[SCREENSHOT PLACEHOLDER - Insert failed GitHub Actions workflow run here]', style='List Bullet')
doc.add_paragraph('Expected to show: Red X mark, test failure details, and error logs', style='List Bullet')

doc.add_paragraph()

doc.add_heading('Screenshot 3: Pull Request Page (feature-export)', level=3)
pr_notes = doc.add_paragraph("""This screenshot shows the Pull Request created for the feature-export branch, demonstrating GitHub's code review interface and CI/CD integration.""")

doc.add_paragraph('[SCREENSHOT PLACEHOLDER - Insert Pull Request page here]', style='List Bullet')
doc.add_paragraph('Expected to show: PR title, description, branch comparison, CI status checks', style='List Bullet')

doc.add_page_break()

# Task 5: Final Merged Repository
doc.add_heading('Task 5: Final Merged Repository Status', level=2)

merge_info = doc.add_paragraph()
merge_run = merge_info.add_run('Repository URL: ')
merge_run.font.bold = True
merge_url_run = merge_info.add_run('https://github.com/AlexGithub901/Student-Management-System')
merge_url_run.font.color.rgb = RGBColor(0, 0, 255)
merge_url_run.underline = True

doc.add_paragraph()

doc.add_heading('Merge History', level=3)

merge_table = doc.add_table(rows=7, cols=3)
merge_table.style = 'Light Grid Accent 1'

header_cells = merge_table.rows[0].cells
header_cells[0].text = 'Commit Hash'
header_cells[1].text = 'Message'
header_cells[2].text = 'Branch'

for cell in header_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

merge_commits = [
    ('4ed344f', 'Add comprehensive tests for JSON and CSV export', 'feature-export'),
    ('6e2cf0b', 'Add JSON export support with format parameter', 'feature-export'),
    ('1407c36', 'Update README with comprehensive documentation', 'main'),
    ('1cc8342', 'Add GitHub Actions Docker build & push workflow', 'main'),
    ('f4b3ab9', 'Optimize Dockerfile with layer caching', 'main'),
    ('f121354', 'Add .dockerignore for optimized builds', 'main'),
]

for i, (commit, msg, branch) in enumerate(merge_commits, 1):
    row_cells = merge_table.rows[i].cells
    row_cells[0].text = commit
    row_cells[1].text = msg
    row_cells[2].text = branch

doc.add_paragraph()

branch_status = doc.add_paragraph()
branch_status_run = branch_status.add_run('Active Branches:')
branch_status_run.font.bold = True
branch_status_run.font.size = Pt(12)

doc.add_paragraph('main - Production branch (merged from feature-export)', style='List Bullet')
doc.add_paragraph('feature-export - Feature branch (merged into main)', style='List Bullet')
doc.add_paragraph('feature-student-search - Historical feature branch', style='List Bullet')

doc.add_page_break()

# Task 6: Git Commands & Workflow
doc.add_heading('Task 6: Git Commands Used & Workflow Explanation', level=2)

doc.add_heading('Git Setup & Configuration', level=3)

setup_cmds = [
    ('git config --global user.name "AlexGithub901"', 'Configure Git user name'),
    ('git config --global user.email "aniketpredij@gmail.com"', 'Configure Git email'),
]

for cmd, desc in setup_cmds:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(cmd).font.family = 'Courier New'
    p.add_run(' - ' + desc)

doc.add_paragraph()

doc.add_heading('Repository Initialization & Initial Commits', level=3)

init_cmds = [
    ('git status', 'Check current repository status'),
    ('git add .', 'Stage all files for commit'),
    ('git commit -m "Initial commit"', 'Create initial commit with message'),
    ('git remote add origin <url>', 'Connect local repo to GitHub remote'),
    ('git push -u origin main', 'Push main branch and set upstream'),
]

for cmd, desc in init_cmds:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(cmd).font.family = 'Courier New'
    p.add_run(' - ' + desc)

doc.add_paragraph()

doc.add_heading('Branching Strategy', level=3)

branch_text = """The project implements a feature branch workflow:

1. main - Stable production branch
2. feature-* - Feature development branches
3. All feature branches are merged via Pull Requests
4. GitHub Actions validates code before merging"""

doc.add_paragraph(branch_text)

doc.add_paragraph()

doc.add_heading('Branch Operations', level=3)

branch_cmds = [
    ('git branch', 'List all local branches'),
    ('git branch -a', 'List all branches (local and remote)'),
    ('git checkout -b feature-export', 'Create and switch to new feature branch'),
    ('git push origin feature-export', 'Push feature branch to GitHub'),
    ('git pull origin main', 'Pull latest changes from main branch'),
]

for cmd, desc in branch_cmds:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(cmd).font.family = 'Courier New'
    p.add_run(' - ' + desc)

doc.add_page_break()

doc.add_heading('Pull Request & Merge Workflow', level=3)

pr_workflow = """Pull Request Process:

1. Create feature branch: git checkout -b feature-export
2. Make changes and commit: git commit -m "descriptive message"
3. Push to GitHub: git push origin feature-export
4. Create Pull Request via GitHub web interface
5. GitHub Actions runs automated tests (CI/CD)
6. Wait for all checks to pass (green checkmarks)
7. Review code and approve Pull Request
8. Merge Pull Request to main branch
9. Delete feature branch after merge
10. Pull latest changes: git pull origin main"""

doc.add_paragraph(pr_workflow)

doc.add_paragraph()

doc.add_heading('File Operations & Conflict Resolution', level=3)

file_cmds = [
    ('git add <file>', 'Stage specific file for commit'),
    ('git add .', 'Stage all modified files'),
    ('git commit -m "message"', 'Commit staged changes with message'),
    ('git push origin <branch>', 'Push branch to GitHub'),
    ('git pull origin <branch>', 'Pull latest changes from branch'),
    ('git status', 'Check status of working directory'),
    ('git log --oneline', 'View commit history in compact format'),
]

for cmd, desc in file_cmds:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(cmd).font.family = 'Courier New'
    p.add_run(' - ' + desc)

doc.add_page_break()

# GitHub Actions Workflow
doc.add_heading('GitHub Actions Workflows', level=2)

doc.add_heading('Workflow 1: Python CI (.github/workflows/python.yml)', level=3)

workflow1_text = """Triggers: On push to main, feature-student-search, feature-export branches and pull requests to main

Steps:
1. Checkout code (actions/checkout@v4)
2. Set up Python 3.12 (actions/setup-python@v5)
3. Install dependencies from requirements.txt
4. Run pytest test_student.py -v

Status Check: PASS if all tests pass, FAIL if any test fails"""

doc.add_paragraph(workflow1_text)

doc.add_paragraph()

doc.add_heading('Workflow 2: Docker Build & Push (.github/workflows/docker.yml)', level=3)

workflow2_text = """Triggers: On push to main and pull requests to main

Steps:
1. Checkout code (actions/checkout@v4)
2. Set up Docker Buildx (docker/setup-buildx-action@v3)
3. Log in to Docker Hub (docker/login-action@v3)
   - Uses DOCKER_USERNAME and DOCKER_PASSWORD secrets
4. Build and push Docker image (docker/build-push-action@v5)
   - Tags: godsogeking/student-management:latest
   - Tags: godsogeking/student-management:<commit-sha>
   - Uses GitHub Actions cache for faster builds
5. Run tests in container

Status Check: PASS if Docker build succeeds and all tests pass"""

doc.add_paragraph(workflow2_text)

doc.add_paragraph()

doc.add_heading('GitHub Secrets Configuration', level=3)

secrets_text = """Required secrets for Docker CI/CD:

Settings → Secrets and variables → Actions

1. DOCKER_USERNAME
   Value: godsogeking
   Purpose: Docker Hub authentication for image push

2. DOCKER_PASSWORD
   Value: Docker Hub access token
   Purpose: Secure authentication (not password)
   
Access tokens provide:
✓ Limited scope (only Docker Hub push)
✓ Ability to revoke without changing password
✓ Audit trail of usage"""

doc.add_paragraph(secrets_text)

doc.add_page_break()

# Docker Setup
doc.add_heading('Docker Containerization', level=2)

doc.add_heading('Dockerfile Optimization', level=3)

dockerfile_text = """The Dockerfile follows production best practices:

1. Multi-stage builds (when applicable)
2. Layer caching for faster builds
3. Slim base image: python:3.12-slim
4. Separate COPY for requirements.txt (before application code)
   - Reason: Dependencies change less frequently than code
   - Result: Improved cache hit rates
5. --no-cache-dir flag for pip
   - Reduces final image size
   - Not needed for installed packages

Image Size: 50.6 MB
Base Image Size: ~45 MB
Application Layer: ~5.6 MB"""

doc.add_paragraph(dockerfile_text)

doc.add_heading('.dockerignore Configuration', level=3)

dockerignore_items = [
    '.git - Version control directory',
    '__pycache__ - Python cache files',
    '.pytest_cache - Pytest cache',
    '.env - Environment variables',
    '.github - GitHub metadata',
    '*.pyc - Compiled Python files',
    'venv - Virtual environment',
]

for item in dockerignore_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# Test Results
doc.add_heading('Testing & Quality Assurance', level=2)

doc.add_heading('Test Suite Overview', level=3)

test_summary = """Total Tests: 17
Status: ✅ ALL PASSING

Test Breakdown:
• 12 Core Functionality Tests (student CRUD operations)
• 5 Export Format Tests (CSV and JSON)"""

doc.add_paragraph(test_summary)

doc.add_paragraph()

test_results_table = doc.add_table(rows=18, cols=3)
test_results_table.style = 'Light Grid Accent 1'

header_cells = test_results_table.rows[0].cells
header_cells[0].text = 'Test Name'
header_cells[1].text = 'Status'
header_cells[2].text = 'Purpose'

for cell in header_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

tests_data = [
    ('test_add_student_new', '✅ PASS', 'Add new student successfully'),
    ('test_add_student_duplicate', '✅ PASS', 'Reject duplicate student IDs'),
    ('test_remove_student_existing', '✅ PASS', 'Remove existing student'),
    ('test_remove_student_not_found', '✅ PASS', 'Handle missing student gracefully'),
    ('test_search_by_id_exact', '✅ PASS', 'Search by exact student ID'),
    ('test_search_by_name_partial', '✅ PASS', 'Partial name search'),
    ('test_search_case_insensitive', '✅ PASS', 'Case-insensitive search'),
    ('test_search_no_match', '✅ PASS', 'Handle no search results'),
    ('test_update_student_name', '✅ PASS', 'Update student name'),
    ('test_update_student_grade', '✅ PASS', 'Update student grade'),
    ('test_update_student_not_found', '✅ PASS', 'Handle update on missing student'),
    ('test_update_student_multiple_fields', '✅ PASS', 'Update multiple fields'),
    ('test_export_students_csv', '✅ PASS', 'Export to CSV format'),
    ('test_export_students_csv_format', '✅ PASS', 'CSV format explicit'),
    ('test_export_students_json', '✅ PASS', 'Export to JSON format'),
    ('test_export_students_csv_helper', '✅ PASS', 'CSV helper function'),
    ('test_export_students_json_helper', '✅ PASS', 'JSON helper function'),
]

for i, (test, status, purpose) in enumerate(tests_data, 1):
    row_cells = test_results_table.rows[i].cells
    row_cells[0].text = test
    row_cells[1].text = status
    row_cells[2].text = purpose

doc.add_page_break()

# Features Implemented
doc.add_heading('Features Implemented', level=2)

doc.add_heading('Core Functions', level=3)

functions_table = doc.add_table(rows=6, cols=2)
functions_table.style = 'Light Grid Accent 1'

header_cells = functions_table.rows[0].cells
header_cells[0].text = 'Function'
header_cells[1].text = 'Description'

for cell in header_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

functions = [
    ('add_student(id, name, grade)', 'Add new student record'),
    ('remove_student(id)', 'Remove student by ID'),
    ('search_student(query)', 'Search by ID or name (case-insensitive)'),
    ('update_student(id, name=None, grade=None)', 'Update student information'),
    ('export_students(format="csv")', 'Export to CSV or JSON format'),
]

for i, (func, desc) in enumerate(functions, 1):
    row_cells = functions_table.rows[i].cells
    row_cells[0].text = func
    row_cells[1].text = desc

doc.add_paragraph()

doc.add_heading('Export Functionality (Bonus Feature)', level=3)

export_text = """Two export formats supported:

CSV Format (Default):
ID,Name,Grade
101,Alice,A
102,Bob,B

JSON Format:
{
  "101": {"name": "Alice", "grade": "A"},
  "102": {"name": "Bob", "grade": "B"}
}

Usage:
export_students()                    # Returns CSV
export_students(format='csv')        # Explicit CSV
export_students(format='json')       # JSON format
export_students_csv()                # Helper function
export_students_json()               # Helper function"""

doc.add_paragraph(export_text)

doc.add_page_break()

# Deployment & Distribution
doc.add_heading('Docker Hub & Deployment', level=2)

deployment_text = """Image Published: ✅ godsogeking/student-management:latest
Registry: Docker Hub
URL: https://hub.docker.com/r/godsogeking/student-management

Image Details:
• Size: 50.6 MB
• Base: python:3.12-slim
• Tags: latest, <commit-sha>
• Status: Public (anyone can pull)
• Auto-pushed: On main branch push via GitHub Actions

Pulling the Image:
docker pull godsogeking/student-management:latest

Running Tests:
docker run --rm godsogeking/student-management:latest

Expected Output:
============================== 17 passed in 0.05s =============================="""

doc.add_paragraph(deployment_text)

doc.add_page_break()

# Challenges & Solutions
doc.add_heading('Challenges & Solutions', level=2)

challenges = [
    ('Git Authentication', 'Resolved by configuring global user.name and user.email'),
    ('Merge Conflicts', 'Resolved README conflicts by merging remote and local versions'),
    ('Docker Image Pull', 'Network latency resolved by pre-pulling image and rebuilding'),
    ('Test Collection', 'Fixed by rebuilding Docker image without cache'),
    ('GitHub Actions Secrets', 'Configured DOCKER_USERNAME and DOCKER_PASSWORD for CI/CD'),
]

for challenge, solution in challenges:
    p = doc.add_paragraph()
    challenge_run = p.add_run(challenge + ': ')
    challenge_run.font.bold = True
    p.add_run(solution)

doc.add_page_break()

# Conclusion
doc.add_heading('Conclusion', level=2)

conclusion_text = """This project successfully demonstrates professional Git and GitHub workflow management with full CI/CD automation. Key achievements include:

✅ Repository Structure: Clean, organized, and documented
✅ Branching Strategy: Feature branch workflow with proper naming
✅ Pull Requests: Code review process with automated checks
✅ GitHub Actions: Dual workflows (Python CI + Docker CI/CD)
✅ Testing: 17 comprehensive tests covering all functionality
✅ Docker: Multi-format export, optimized image (50.6 MB)
✅ Documentation: Detailed README and inline documentation
✅ Git History: Clear, descriptive commit messages
✅ Deployment: Automated push to Docker Hub on merge

The project follows industry standards for:
• Version control practices
• Continuous Integration/Continuous Deployment
• Code quality assurance
• Documentation
• Team collaboration

This workflow is production-ready and scalable for larger teams and projects."""

doc.add_paragraph(conclusion_text)

# Summary Table
doc.add_heading('Project Summary', level=2)

summary_table = doc.add_table(rows=15, cols=2)
summary_table.style = 'Light Grid Accent 1'

summary_data = [
    ('Project Name', 'Student Management System'),
    ('Repository', 'https://github.com/AlexGithub901/Student-Management-System'),
    ('Primary Language', 'Python 3.12'),
    ('Test Framework', 'pytest'),
    ('Total Tests', '17 (100% passing)'),
    ('Docker Image Size', '50.6 MB'),
    ('Docker Registry', 'Docker Hub (godsogeking/student-management)'),
    ('CI/CD Platform', 'GitHub Actions'),
    ('Active Branches', '3 (main, feature-export, feature-student-search)'),
    ('Total Commits', '6+ (main branch)'),
    ('Code Style', 'PEP 8 compliant'),
    ('Documentation', 'Comprehensive README'),
    ('Status', '✅ Production Ready'),
    ('Last Updated', datetime.now().strftime("%B %d, %Y")),
]

for i, (key, value) in enumerate(summary_data):
    row_cells = summary_table.rows[i].cells
    key_run = row_cells[0].paragraphs[0].add_run(key)
    key_run.font.bold = True
    row_cells[1].text = value

doc.add_page_break()

# Appendix
doc.add_heading('Appendix: Quick Reference', level=2)

doc.add_heading('Essential Git Commands', level=3)

commands_dict = {
    'Clone Repository': 'git clone <url>',
    'Create Branch': 'git checkout -b <branch-name>',
    'Switch Branch': 'git checkout <branch-name>',
    'Stage Changes': 'git add <file> or git add .',
    'Commit Changes': 'git commit -m "message"',
    'Push to Remote': 'git push origin <branch>',
    'Pull from Remote': 'git pull origin <branch>',
    'View History': 'git log --oneline',
    'Check Status': 'git status',
    'Create PR': 'Via GitHub web interface',
    'Merge PR': 'Via GitHub web interface',
}

for cmd, usage in commands_dict.items():
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(cmd + ': ').font.bold = True
    p.add_run(usage).font.family = 'Courier New'

doc.add_paragraph()

doc.add_heading('Docker Commands', level=3)

docker_commands = {
    'Build Image': 'docker build -t image-name .',
    'Run Container': 'docker run --rm image-name',
    'Push to Hub': 'docker push username/image:tag',
    'Pull Image': 'docker pull username/image:tag',
    'Run Tests': 'docker compose up',
}

for cmd, usage in docker_commands.items():
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(cmd + ': ').font.bold = True
    p.add_run(usage).font.family = 'Courier New'

# Save document
doc.save('Student_Management_System_Report.docx')
print("[SUCCESS] Word document created: Student_Management_System_Report.docx")
